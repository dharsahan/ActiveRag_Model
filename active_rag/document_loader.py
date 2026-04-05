"""Document ingestion for local files (TXT, MD, PDF, DOCX)."""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from .config import Config

logger = logging.getLogger(__name__)

_SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    """A document loaded from a local file."""
    content: str
    source: str
    title: str = ""
    word_count: int = 0


class DocumentLoader:
    """Loads documents from local files for vector store ingestion."""

    def __init__(self, config: Optional[Config] = None):
        """Initialize DocumentLoader"""
        self.config = config or Config()

    def load(self, path: str) -> List[LoadedDocument]:
        """Load a document from *path* and return parsed content."""
        p = Path(path).expanduser().resolve()
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = p.suffix.lower()
        if ext not in _SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
            )

        if ext == ".txt":
            return self._load_text(p)
        elif ext == ".md":
            return self._load_markdown(p)
        elif ext == ".pdf":
            return self._load_pdf(p)
        elif ext == ".docx":
            return self._load_docx(p)
        return []

    def _load_text(self, path: Path) -> list[LoadedDocument]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [LoadedDocument(
            content=text, source=str(path),
            title=path.stem, word_count=len(text.split()),
        )]

    def _load_markdown(self, path: Path) -> list[LoadedDocument]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        # Strip markdown formatting
        text = re.sub(r"#{1,6}\s*", "", text)  # headers
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
        text = re.sub(r"\*(.+?)\*", r"\1", text)  # italic
        text = re.sub(r"`(.+?)`", r"\1", text)  # inline code
        return [LoadedDocument(
            content=text, source=str(path),
            title=path.stem, word_count=len(text.split()),
        )]

    def _load_pdf(self, path: Path) -> list[LoadedDocument]:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        pages: list[LoadedDocument] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(LoadedDocument(
                    content=text, source=f"{path}#page={i+1}",
                    title=f"{path.stem} (p{i+1})", word_count=len(text.split()),
                ))
        return pages

    def _load_docx(self, path: Path) -> list[LoadedDocument]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [LoadedDocument(
            content=text, source=str(path),
            title=path.stem, word_count=len(text.split()),
        )]
